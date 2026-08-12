from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION

OUT = r"C:\Users\Javier\Desktop\sentinel-taskmaster\Sentinel_Taskmaster_Especificacion_Tecnica_ES_Corregida.docx"

INK = "17212B"
CHARCOAL = "2B3540"
MUTED = "66717D"
LIGHT = "F2F4F6"
MID = "D9DEE3"
WHITE = "FFFFFF"
CAUTION = "5D4A00"
RISK = "7D2525"


def set_font(run, size=None, bold=None, color=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = borders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            spec = kwargs[edge]
            element.set(qn("w:val"), spec.get("val", "single"))
            element.set(qn("w:sz"), str(spec.get("sz", 4)))
            element.set(qn("w:color"), spec.get("color", MID))


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "140")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        grid.gridCol_lst[idx].set(qn("w:w"), str(width))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True):
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for cell in row.cells:
            set_cell_border(cell,
                top={"color": MID, "sz": 5}, bottom={"color": MID, "sz": 5},
                left={"color": MID, "sz": 5}, right={"color": MID, "sz": 5})
            if header and r_idx == 0:
                shade(cell, CHARCOAL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_font(run, 9.2, bold=(header and r_idx == 0), color=(WHITE if header and r_idx == 0 else INK))


def add_run(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r


def add_body(doc, text, after=6, bold_prefix=None):
    p = doc.add_paragraph(style="Body")
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, size=10.6, bold=True, color=INK)
        add_run(p, text[len(bold_prefix):], size=10.6, color=INK)
    else:
        add_run(p, text, size=10.6, color=INK)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_run(p, text, size={1: 16, 2: 12.5, 3: 11.2}[level], bold=True, color=INK)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, text, size=10.3, color=INK)


def add_callout(doc, label, text, color=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, color)
    set_cell_border(cell, top={"color": MID, "sz": 6}, bottom={"color": MID, "sz": 6}, left={"color": MID, "sz": 6}, right={"color": MID, "sz": 6})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"{label}  ", size=10.2, bold=True, color=INK)
    add_run(p, text, size=10.2, color=INK)
    style_table(table, header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, size=8.3, color=MUTED, italic=True)


def add_page_number(p):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.12
    for lvl, before, after in ((1, 16, 8), (2, 12, 6), (3, 8, 4)):
        s = styles[f"Heading {lvl}"]
        s.font.name = "Aptos Display"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        s.font.color.rgb = RGBColor.from_string(INK)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "SENTINEL TASKMASTER  |  ESPECIFICACIÓN TÉCNICA", size=8.5, bold=True, color=MUTED)
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "All Things Agentic Hackathon  •  ", size=8.5, color=MUTED)
    add_page_number(p)


def title_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "ALL THINGS AGENTIC HACKATHON  •  TASKMASTER", size=10, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "SENTINEL\nTASKMASTER", size=31, bold=True, color=INK, name="Aptos Display")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    add_run(p, "Especificación técnica y plan de demostración", size=15, color=CHARCOAL)
    add_callout(doc, "Tesis de innovación", "Autonomía controlada: un agente que resuelve un incidente operativo, actúa dentro de límites verificables y solicita intervención humana cuando la autoridad autónoma no es suficiente.")
    tbl = doc.add_table(rows=4, cols=2)
    set_table_geometry(tbl, [2500, 6860])
    rows = [
        ("Estado", "Línea base corregida para implementación"),
        ("Stack", "Gemini 3.5+ • Google ADK • Python • Cloud Run"),
        ("Entrega", "Aplicación, repositorio, diagrama, README, vídeo y Devpost"),
        ("Fecha límite", "31 de agosto de 2026, 5:00 p. m. PDT"),
    ]
    for row, (a, b) in zip(tbl.rows, rows):
        row.cells[0].text = a
        row.cells[1].text = b
        shade(row.cells[0], LIGHT)
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, 9.6, bold=True, color=INK)
        for run in row.cells[1].paragraphs[0].runs:
            set_font(run, 9.6, color=INK)
    style_table(tbl, header=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Documento de trabajo • Versión corregida • 12 de agosto de 2026", size=9, color=MUTED)
    doc.add_page_break()


def executive(doc):
    add_heading(doc, "1. Resumen ejecutivo")
    add_body(doc, "Sentinel Taskmaster es un agente autónomo de respuesta a incidentes operativos. Recibe un objetivo estructurado, investiga evidencia, formula un plan, opera herramientas autorizadas y verifica de manera independiente que el estado final sea correcto.")
    add_body(doc, "El diferenciador no es que el agente “converse”, sino que ejecuta un flujo completo con gobierno visible: evalúa riesgo y presupuesto antes de actuar, valida el origen de la memoria, conserva trazabilidad y solicita aprobación humana cuando una acción supera su autoridad.")
    add_callout(doc, "Decisión de producto", "La primera demo se construirá sobre un entorno simulado, determinista y seguro. El caso de uso debe representar una fricción real y repetible: recuperar un servicio de pedidos tras una degradación operativa.")
    add_heading(doc, "Resultado esperado de la demo", 2)
    add_bullets(doc, [
        "El agente identifica la causa probable de la degradación a partir de logs, estado del servicio y evidencia contextual.",
        "Ejecuta acciones de bajo riesgo por sí mismo y bloquea una acción de alto riesgo hasta obtener una aprobación explícita.",
        "Rechaza o pone en cuarentena información engañosa antes de incorporarla a la memoria de trabajo.",
        "Verifica de forma independiente que el servicio cumple los criterios de recuperación y registra la trayectoria completa.",
    ])
    add_heading(doc, "Objetivos de éxito", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2350, 3740, 3270])
    for cell, text in zip(table.rows[0].cells, ["Objetivo", "Prueba visible", "Métrica de aceptación"]):
        cell.text = text
    rows = [
        ("Autonomía útil", "Plan + acciones reales sobre el simulador", "El incidente avanza sin intervención en acciones permitidas"),
        ("Control", "Riesgo, política y aprobación humana", "Ninguna acción alta se ejecuta sin autorización"),
        ("Confiabilidad", "Verificador y replanificación", "El cierre requiere evidencia, no solo una respuesta del modelo"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (a, b, c)):
            cell.text = text
    style_table(table)
    add_caption(doc, "Tabla 1. Objetivos que conectan la experiencia de demo con los criterios de evaluación.")


def alignment(doc):
    add_heading(doc, "2. Alineación con el hackathon")
    add_body(doc, "Sentinel Taskmaster se inscribe en el track Taskmaster: debe demostrar un flujo de trabajo completo que toma acciones y elimina fricción, no un asistente que solo genera texto.")
    add_heading(doc, "Requisitos de participación aplicados", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2800, 3900, 2660])
    for cell, text in zip(table.rows[0].cells, ["Regla actual", "Respuesta de Sentinel", "Estado"]): cell.text = text
    rows = [
        ("Gemini 3.5 o superior mediante API o Vertex AI", "Usar Gemini 3.5 Flash o superior; fijar la versión exacta en el repositorio.", "Pendiente de fijar"),
        ("Un framework de agentes de Google", "Google ADK como framework principal.", "Seleccionado"),
        ("Un servicio de infraestructura Google Cloud", "Cloud Run para backend; Firestore como memoria persistente candidata.", "Seleccionado / validar"),
        ("Flujo completo que toma acciones", "Investigación, planificación, ejecución, verificación y replanificación sobre un simulador.", "Diseñado"),
        ("Prueba de despliegue en Google Cloud", "Cloud Run visible en vídeo y documentado en el repositorio.", "Requisito de entrega"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_caption(doc, "Tabla 2. Matriz de cumplimiento. Fuente: página oficial del All Things Agentic Hackathon, consultada el 12 de agosto de 2026.")
    add_heading(doc, "Cómo se evaluará", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2600, 1240, 5520])
    for cell, text in zip(table.rows[0].cells, ["Criterio", "Peso", "Evidencia propuesta"]): cell.text = text
    for values in [
        ("Innovación y utilidad operativa", "40%", "Incidente realista, acciones autónomas y reducción visible de fricción."),
        ("Disciplina arquitectónica y stack", "30%", "Capas desacopladas, estado/memoria segura, credenciales y fallos controlados."),
        ("Demo y preparación para producción", "30%", "Vídeo sin cortes, repositorio reproducible, diagrama y prueba de Cloud Run."),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_callout(doc, "Implicación práctica", "No conviene optimizar solo la arquitectura. La demo debe hacer visible qué decidió el agente, qué acción ejecutó, qué bloqueó y con qué evidencia verificó el resultado.")


def mission(doc):
    add_heading(doc, "3. Misión operativa de la demo")
    add_body(doc, "Caso de uso recomendado: un servicio simulado de ingestión de pedidos presenta latencia elevada y una cola acumulada. El agente debe recuperar el servicio respetando un presupuesto de acciones y una política que protege cambios de alto impacto.")
    add_heading(doc, "Contrato de misión", 2)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    for cell, text in zip(table.rows[0].cells, ["Elemento", "Definición"]): cell.text = text
    for values in [
        ("Objetivo", "Restablecer la tasa de procesamiento y reducir la cola por debajo del umbral definido."),
        ("Entorno", "Simulador con estado del servicio, cola de pedidos, logs, activos, políticas y catálogo de acciones."),
        ("Riesgo bajo/medio", "Consultar logs, diagnosticar estado, reiniciar un worker aislado, limpiar un lote corrupto reversible."),
        ("Riesgo alto", "Aumentar capacidad global o vaciar una cola de recuperación; requiere aprobación humana."),
        ("Entrada no confiable", "Un log inyectado propone una solución destructiva; debe quedar en cuarentena por procedencia insuficiente."),
        ("Criterio de cierre", "Latencia y cola dentro de umbral, no hay alertas críticas y el verificador confirma consistencia."),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_heading(doc, "Ciclo de ejecución", 2)
    flow = doc.add_table(rows=1, cols=7)
    set_table_geometry(flow, [1220, 180, 1220, 180, 1220, 180, 5160])
    texts = ["INCIDENTE", "→", "EVIDENCIA", "→", "PLAN", "→", "POLÍTICA • EJECUCIÓN • VERIFICACIÓN"]
    for cell, text in zip(flow.rows[0].cells, texts):
        cell.text = text
        shade(cell, CHARCOAL if text != "→" else WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: set_font(r, 8.2, bold=(text != "→"), color=(WHITE if text != "→" else MUTED))
    style_table(flow, header=False)
    add_callout(doc, "Bucle de control", "Tras cada acción: observar el cambio de estado, actualizar el plan, confirmar que el presupuesto sigue disponible y solo cerrar cuando el verificador independiente produzca evidencia suficiente.")


def controls(doc):
    add_heading(doc, "4. Autonomía controlada")
    add_body(doc, "La política no es una instrucción opcional dentro del prompt. Es una capa de control independiente que recibe la acción propuesta, la evidencia disponible, el nivel de riesgo, el presupuesto restante y el estado de aprobación humana.")
    add_heading(doc, "Decisión por acción", 2)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2200, 2100, 2470, 2590])
    for cell, text in zip(table.rows[0].cells, ["Entrada", "Evaluación", "Salida", "Registro mínimo"]): cell.text = text
    for values in [
        ("Acción propuesta por planner", "Riesgo + política + presupuesto", "Permitir, bloquear o solicitar aprobación", "Razón, política aplicada y nivel de confianza"),
        ("Dato candidato a memoria", "Procedencia + consistencia + relevancia", "Guardar, rechazar o cuarentena", "Fuente, evidencia y estado de validación"),
        ("Resultado de herramienta", "Integridad + efecto en el objetivo", "Continuar, replanificar o fallar", "Observación, cambio de plan y métrica"),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_heading(doc, "Política inicial", 2)
    add_bullets(doc, [
        "Nunca ejecutar una acción marcada como alta sin aprobación humana registrada y vigente.",
        "No aceptar hechos en memoria sin fuente trazable o evidencia corroborante.",
        "Detener y replanificar si una acción no genera el efecto esperado o si se excede el presupuesto permitido.",
        "Separar el verificador del planificador: el mismo razonamiento no debe auto-certificar el éxito sin datos del entorno.",
    ])
    add_callout(doc, "Aprobación humana", "La interfaz debe mostrar: acción solicitada, impacto esperado, riesgo, alternativa reversible, evidencia y botones Aprobar / Rechazar. La decisión se incorpora a la trayectoria de auditoría.", "FFF7E6")


def architecture(doc):
    add_heading(doc, "5. Arquitectura objetivo")
    add_body(doc, "La arquitectura se mantiene deliberadamente pequeña. Cada componente existe para resolver una responsabilidad visible en la demo y en el repositorio; no se añaden servicios de Google Cloud sin una función verificable.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2680, 3300, 3380])
    for cell, text in zip(table.rows[0].cells, ["Capa", "Responsabilidad", "Implementación inicial"]): cell.text = text
    rows = [
        ("Interfaz de incidente", "Recibe misión, muestra estado y captura aprobación humana.", "UI web mínima; backend HTTP en Cloud Run."),
        ("Orquestación ADK", "Coordina planner, ejecutor, memoria y verificador.", "Python + Google ADK + Gemini 3.5+."),
        ("Policy & risk engine", "Clasifica acciones y aplica límites antes de ejecutar.", "Módulo determinista, separado del prompt."),
        ("Herramientas del simulador", "Lee y modifica el entorno bajo contratos definidos.", "APIs locales simuladas con respuestas trazables."),
        ("Memoria y auditoría", "Conserva estado persistente y trayectoria de misión.", "Firestore candidato; logs estructurados en Cloud Logging cuando aporte valor."),
        ("Verificador", "Comprueba criterios de éxito con datos del entorno.", "Checks deterministas + resumen de evidencia."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_heading(doc, "Principios de implementación", 2)
    add_bullets(doc, [
        "Las herramientas exponen contratos estrechos: entrada validada, salida estructurada y efectos auditablemente reversibles cuando sea posible.",
        "Los secretos y credenciales no se entregan al modelo; se resuelven en la capa de ejecución autorizada.",
        "El estado del simulador, las decisiones de política y los resultados de herramientas se registran como eventos, no como texto libre solamente.",
        "El despliegue debe probarse en Cloud Run y verse en el vídeo o en la consola de Google Cloud.",
    ])


def functional(doc):
    add_heading(doc, "6. Requisitos funcionales y no funcionales")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1140, 2630, 5590])
    for cell, text in zip(table.rows[0].cells, ["ID", "Capacidad", "Criterio de aceptación"]): cell.text = text
    rows = [
        ("RF-01", "Recepción de misión", "Acepta un incidente estructurado con objetivo, umbrales y restricciones."),
        ("RF-02", "Planificación", "Genera pasos ordenados vinculados a evidencia y herramientas disponibles."),
        ("RF-03", "Ejecución", "Llama herramientas del entorno y captura resultados estructurados."),
        ("RF-04", "Riesgo y política", "Clasifica cada acción y evita que el modelo eluda restricciones."),
        ("RF-05", "Presupuesto", "Mide y limita las acciones o recursos permitidos por misión."),
        ("RF-06", "Memoria segura", "Guarda procedencia y permite rechazar o poner datos en cuarentena."),
        ("RF-07", "Aprobación humana", "Pausa para acciones altas y registra la decisión humana."),
        ("RF-08", "Verificación y replanificación", "Cierra solo con evidencia independiente o crea un plan alternativo."),
        ("RF-09", "Auditoría", "Registra intención, plan, decisión de política, acción, observación y resultado."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_caption(doc, "Tabla 3. Requisitos priorizados para un MVP evaluable.")
    add_heading(doc, "Requisitos no funcionales", 2)
    add_bullets(doc, [
        "Determinismo de demo: la misma misión debe producir estados observables y reproducibles.",
        "Seguridad: las acciones se limitan al entorno simulado y a una política explícita.",
        "Observabilidad: cada acción debe ser reconstruible desde la trayectoria de misión.",
        "Economía: el diseño debe poder ejecutarse con costes bajos y apagarse fuera de la demostración.",
    ])


def demo(doc):
    add_heading(doc, "7. Guion de demostración (4 minutos)")
    add_body(doc, "La demo debe ser una historia única, directa y sin cortes: problema realista → razonamiento visible → acción → control humano → recuperación verificada. Evita presentar pantallas estáticas o explicación prolongada del código.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1150, 2900, 5310])
    for cell, text in zip(table.rows[0].cells, ["Tiempo", "Escena", "Prueba que debe verse"]): cell.text = text
    rows = [
        ("0:00", "Incidente y objetivo", "Panel con latencia, cola y umbrales de éxito."),
        ("0:25", "Plan inicial", "Pasos propuestos, evidencia inicial y presupuesto."),
        ("0:55", "Investigación", "Herramientas consultan logs/estado; se muestra una entrada sospechosa."),
        ("1:30", "Memoria y política", "Dato engañoso en cuarentena; acción baja permitida con razón."),
        ("2:05", "Acción de alto riesgo", "Solicitud de aprobación humana con impacto y alternativa."),
        ("2:30", "Replanificación", "Resultado parcial obliga a ajustar el plan y ejecutar la siguiente acción."),
        ("3:15", "Verificación", "Métricas dentro de umbral y comprobación independiente."),
        ("3:40", "Prueba de Cloud", "Cloud Run / consola / URL y arquitectura resumida."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_callout(doc, "Regla de oro", "Cada 20–30 segundos el espectador debe ver una consecuencia: evidencia nueva, una decisión, un cambio en el entorno o una verificación. La voz explica; la interfaz demuestra.")


def plan(doc):
    add_heading(doc, "8. Plan de implementación")
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [780, 2350, 3960, 2270])
    for cell, text in zip(table.rows[0].cells, ["Fase", "Resultado", "Trabajo", "Salida verificable"]): cell.text = text
    rows = [
        ("1", "Diseño de misión", "Definir estado, herramientas, políticas, umbrales y datos engañosos.", "Fixture del incidente y contrato del simulador."),
        ("2", "MVP de agente", "Configurar ADK, Gemini y planner; implementar el estado de misión.", "Plan y herramienta de consulta funcionando."),
        ("3", "Control", "Implementar riesgo, presupuesto, memoria validada y aprobación humana.", "Acción alta bloqueada y aprobable."),
        ("4", "Cierre del ciclo", "Añadir verificador, replanificación y auditoría estructurada.", "Misión completa con evidencia de éxito."),
        ("5", "Cloud y entrega", "Desplegar en Cloud Run, documentar y grabar vídeo.", "Repo reproducible + prueba de despliegue."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_heading(doc, "Decisiones que deben cerrarse antes de la Fase 2", 2)
    add_bullets(doc, [
        "Versión definitiva de Gemini y ruta de acceso (Gemini API o Vertex AI).",
        "Esquema del incidente, umbrales de éxito y transición de estados del simulador.",
        "Lista de herramientas y clasificación de riesgo por herramienta/acción.",
        "Modelo de presupuesto, mecanismo de aprobación y persistencia de memoria/auditoría.",
    ])


def delivery(doc):
    add_heading(doc, "9. Entregables y lista de verificación")
    add_body(doc, "La entrega debe demostrar construcción real, no solo intención. Este checklist reúne los elementos de la especificación y los requisitos de Devpost para reducir el riesgo de una entrega incompleta.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [5400, 2040, 1920])
    for cell, text in zip(table.rows[0].cells, ["Entregable", "Evidencia mínima", "Estado"]): cell.text = text
    rows = [
        ("Aplicación Sentinel Taskmaster", "Flujo completo contra el incidente simulado.", "Por implementar"),
        ("Repositorio GitHub/GitLab/Bitbucket", "Código, dependencias, arquitectura y README reproducible.", "Por crear"),
        ("Instrucciones de arranque", "Paso a paso local y de despliegue; variables y requisitos claros.", "Por redactar"),
        ("Diagrama de arquitectura", "Gemini, ADK, backend, memoria, herramientas y verificador.", "Por crear"),
        ("Vídeo ~4 min", "Problema, valor, ejecución sin cortes y evidencia de Cloud Run.", "Por grabar"),
        ("Project Story Devpost", "Descripción, funcionalidades, tecnologías, fuentes de datos y aprendizajes.", "Por completar"),
        ("Prueba de Google Cloud", "Cloud Run / consola / logs / URL documentados en vídeo y repo.", "Por validar"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values): cell.text = text
    style_table(table)
    add_heading(doc, "Bonos opcionales", 2)
    add_bullets(doc, [
        "Publicar contenido público sobre cómo se construyó el proyecto, indicando que se creó para este hackathon.",
        "Publicar una pieza en redes sociales con el hashtag #AllThingsAgenticHackathon.",
        "Integrar de forma justificada un modelo de Google AI adicional, como Gemma, Veo o Lyria.",
    ])
    add_callout(doc, "Siguiente paso inmediato", "Diseñar el fixture del incidente: estado inicial, evidencia válida y engañosa, catálogo de acciones, efectos esperados, reglas de riesgo y verificador. Ese artefacto desbloquea la implementación completa.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Anexo A. Fuente y control de cambios")
    add_body(doc, "Este documento reemplaza la línea base anterior con una versión corregida y rediseñada. Se actualizaron los requisitos del hackathon, se concretó la misión recomendada, se reforzó la estrategia de evaluación y se reorganizó el contenido para lectura ejecutiva y de implementación.")
    add_body(doc, "Fuente externa: All Things Agentic Hackathon, Devpost (https://allthingsagentichackathon.devpost.com/), consultada el 12 de agosto de 2026.")


def main():
    doc = Document()
    setup(doc)
    title_page(doc)
    executive(doc)
    alignment(doc)
    mission(doc)
    controls(doc)
    architecture(doc)
    functional(doc)
    demo(doc)
    plan(doc)
    delivery(doc)
    doc.core_properties.title = "Sentinel Taskmaster — Especificación Técnica"
    doc.core_properties.subject = "All Things Agentic Hackathon — Track Taskmaster"
    doc.core_properties.author = "Sentinel Taskmaster"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
